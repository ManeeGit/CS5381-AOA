"""Generate CS5381_Presentation_Script.docx from the speaker script content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

C_BLUE  = RGBColor(0x1F, 0x65, 0xAD)
C_BLUE2 = RGBColor(0x2E, 0x74, 0xB5)
C_GRAY  = RGBColor(0x70, 0x70, 0x70)


def heading1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = C_BLUE
    return p


def para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def speaker(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("Speaker: " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_GRAY


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── Title block ──────────────────────────────────────────────────────────────
t = doc.add_heading("CS5381 AOA - Group 2 Presentation Script", 0)
for run in t.runs:
    run.font.color.rgb = C_BLUE

p = doc.add_paragraph()
run = p.add_run('"Evolve" - AlphaEvolve-Inspired Evolutionary Code Optimizer')
run.italic = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run("Target: ~2 minutes (~20 seconds per slide)  |  6 Slides")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = C_BLUE2

doc.add_paragraph()

# ─── Slide 1 ──────────────────────────────────────────────────────────────────
heading1("Slide 1 - Title  (~15 seconds)")
speaker("any member")
para(
    "Hello everyone. Our project is called Evolve - an AlphaEvolve-inspired framework "
    "that uses evolutionary algorithms to automatically improve code. We are Group 2. "
    "Today I will walk you through what we built, how it works, and what the results "
    "showed - in about two minutes."
)
doc.add_paragraph()

# ─── Slide 2 ──────────────────────────────────────────────────────────────────
heading1("Slide 2 - Group Members & Configurations  (~20 seconds)")
speaker("any member")
para(
    "Our team of five each ran the system with a different configuration - different LLM "
    "backends, population sizes, and generation counts - so we could compare results "
    "across machines and settings. All configs use the same seed=42 for reproducibility."
)
doc.add_paragraph()
add_table(
    ["Member", "Config Highlights"],
    [
        ["Bhanu Sankar Ravi",  "Ollama local, matrix + pacman"],
        ["Johitha Konduru",    "Remote LLM, pseudocode"],
        ["Lakshman Pukhraj",   "No-LLM baseline"],
        ["Maneesh Malepati",   "llm_guided, pacman"],
        ["Demo Student",       "All three problems, default config"],
    ],
)
doc.add_paragraph()

# ─── Slide 3 ──────────────────────────────────────────────────────────────────
heading1("Slide 3 - What We Built  (~25 seconds)")
speaker("whoever handles the architecture")
para(
    "We support three problems: Pacman (game-playing agent), Matrix multiplication "
    "(correctness + efficiency), and Pseudocode (sorting algorithm quality with four "
    "configurable fitness dimensions)."
)
para(
    "We have three evolution modes - a no-evolution baseline, random mutation, and "
    "LLM-guided mutation using Ollama with qwen2.5-coder."
)
para("And three interfaces: a live Streamlit UI, a FastAPI REST API, and a CLI profiler.")
doc.add_paragraph()
para("Key fitness functions:", bold=True)
bullet("Pacman:      F = 0.6 * score + 0.3 * survival - 0.1 * steps")
bullet("Matrix:      F = 0.7 * correct + 0.3 * (1 - ops_ratio)")
bullet("Pseudocode:  F = w1*correct + w2*runtime + w3*length + w4*readability")
doc.add_paragraph()

# ─── Slide 4 ──────────────────────────────────────────────────────────────────
heading1("Slide 4 - The Evolutionary Loop  (~20 seconds)")
speaker("whoever handles the algorithm")
para("The loop is five steps - straight from AOA:")
bullet("Initialise   - seed population from a base template")
bullet("Evaluate     - run each candidate through the fitness function")
bullet("Select       - greedy top-K elite selection")
bullet("Mutate       - random perturbation, or LLM-guided rewrite for 20% of candidates")
bullet("Cache & Repeat - vector cache avoids re-evaluating identical or near-identical code")
doc.add_paragraph()
para(
    "This repeats for however many generations you configure. Greedy, randomised, and "
    "approximation - all three AOA paradigms in one loop."
)
doc.add_paragraph()

# ─── Slide 5 ──────────────────────────────────────────────────────────────────
heading1("Slide 5 - Results  (~25 seconds)")
speaker("whoever ran the experiments")
para(
    "Matrix fitness hits near-ceiling at 0.99 across all modes - the seed template was "
    "already close to optimal for a 2x2 problem."
)
para(
    "Pacman is where we see the real difference: LLM-guided mutation reaches +20-45% above "
    "baseline, random mutation gives +8%. That is not luck - it is consistent across runs."
)
para(
    "The caching system cut repeat evaluation cost by 60-80% in later generations. And the "
    "pseudocode evaluator showed bubble sort evolving toward quicksort-like patterns, "
    "fitness going from 0.855 to 0.873."
)
doc.add_paragraph()
add_table(
    ["Mode", "Matrix Fitness", "Pacman Fitness", "vs Baseline"],
    [
        ["no_evolution",    "0.99", "~58",  "-"],
        ["random_mutation", "0.99", "~63",  "+8%"],
        ["llm_guided",      "0.99", "~70+", "+20-45%"],
    ],
)
doc.add_paragraph()

# ─── Slide 6 ──────────────────────────────────────────────────────────────────
heading1("Slide 6 - Demo & Takeaways  (~15 seconds)")
speaker("demo runner")
para("Run command:  streamlit run app.py  ->  http://localhost:8501")
para(
    "Select a problem, tune generations and population size, hit Run - and watch the "
    "fitness bars update in real time with a code diff."
)
doc.add_paragraph()
para("Key Takeaways:", bold=True)
bullet("LLM guidance is consistently better than random - not just luck")
bullet("Greedy top-K selection is simple but effective")
bullet("Modular design: swap evaluator or LLM backend without touching the engine")
bullet("Reproducibility: seed=42 applied to both random + numpy")
bullet("AOA Connections: Greedy (top-K),  Randomised (mutations),  Approximation (fitness budget)")
doc.add_paragraph()
para("GitHub:     github.com/ManeeGit/CS5381-AOA")
para("Demo Video: youtu.be/b2GBt0VEED0")
doc.add_paragraph()

# ─── References ───────────────────────────────────────────────────────────────
heading1("References")
refs = [
    "[1] A. Novikov et al., AlphaEvolve: A coding agent for scientific and algorithmic "
    "discovery, arXiv:2506.13131, Jun. 2025.",
    "[2] S. Tamilselvi, Introduction to Evolutionary Algorithms, IntechOpen, 2022.",
    "[3] UC Berkeley, Introduction to AI - Pacman Projects. http://ai.berkeley.edu",
    "[4] Meta AI, Ollama: Run large language models locally. https://ollama.com",
    "[5] GitHub Repository: https://github.com/ManeeGit/CS5381-AOA",
    "[6] Demo Video: https://youtu.be/b2GBt0VEED0",
]
for r in refs:
    bullet(r)

# ─── Save ─────────────────────────────────────────────────────────────────────
OUT = "CS5381_Presentation_Script.docx"
doc.save(OUT)
print("[OK] Saved:", OUT)
